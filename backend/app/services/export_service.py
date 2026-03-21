"""Export workspace artifacts to DOCX and PDF formats."""

import io
import logging
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    HRFlowable,
    ListFlowable,
    ListItem,
)
from reportlab.lib.colors import HexColor

logger = logging.getLogger(__name__)


def _strip_markdown(text: str) -> str:
    """Minimal strip of markdown syntax for plain-text contexts."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _md_to_rl_markup(text: str) -> str:
    """Convert markdown bold/italic to ReportLab XML markup."""
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"__(.+?)__", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"_(.+?)_", r"<i>\1</i>", text)
    text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
    return text


def export_to_pdf(md_content: str) -> bytes:
    """Convert markdown content to a PDF byte stream using ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    # Custom styles for resume
    styles.add(ParagraphStyle(
        "ResumeH1", parent=styles["Heading1"],
        fontSize=16, spaceAfter=4, textColor=HexColor("#111111"),
    ))
    styles.add(ParagraphStyle(
        "ResumeH2", parent=styles["Heading2"],
        fontSize=13, spaceBefore=12, spaceAfter=4,
        textColor=HexColor("#333333"), borderWidth=0,
    ))
    styles.add(ParagraphStyle(
        "ResumeH3", parent=styles["Heading3"],
        fontSize=11, spaceBefore=8, spaceAfter=3,
        textColor=HexColor("#444444"),
    ))
    styles.add(ParagraphStyle(
        "ResumeBody", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=4,
        textColor=HexColor("#222222"),
    ))
    styles.add(ParagraphStyle(
        "ResumeBullet", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=3,
        leftIndent=18, textColor=HexColor("#222222"),
    ))

    flowables = []
    lines = md_content.split("\n")

    for line in lines:
        line = line.rstrip()

        if not line:
            flowables.append(Spacer(1, 6))
            continue

        # Headings
        if line.startswith("### "):
            flowables.append(Paragraph(_md_to_rl_markup(line[4:]), styles["ResumeH3"]))
            continue
        if line.startswith("## "):
            flowables.append(Paragraph(_md_to_rl_markup(line[3:]), styles["ResumeH2"]))
            flowables.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc")))
            continue
        if line.startswith("# "):
            flowables.append(Paragraph(_md_to_rl_markup(line[2:]), styles["ResumeH1"]))
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            flowables.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#dddddd")))
            continue

        # Bullet points
        if line.lstrip().startswith(("- ", "* ", "\u2022 ")):
            text = re.sub(r"^[\s]*[-*\u2022]\s+", "", line)
            flowables.append(
                Paragraph(f"\u2022  {_md_to_rl_markup(text)}", styles["ResumeBullet"])
            )
            continue

        # Numbered list
        if re.match(r"^\s*(\d+)\.\s", line):
            match = re.match(r"^\s*(\d+)\.\s+(.*)", line)
            if match:
                num, text = match.groups()
                flowables.append(
                    Paragraph(f"{num}.  {_md_to_rl_markup(text)}", styles["ResumeBullet"])
                )
            continue

        # Regular paragraph
        flowables.append(Paragraph(_md_to_rl_markup(line), styles["ResumeBody"]))

    doc.build(flowables)
    return buffer.getvalue()


def export_to_docx(md_content: str) -> bytes:
    """Convert markdown content to a DOCX byte stream."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = md_content.split("\n")
    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(_strip_markdown(line[4:]), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(_strip_markdown(line[3:]), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(_strip_markdown(line[2:]), level=1)
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
            continue

        # Bullet points
        if line.lstrip().startswith(("- ", "* ", "\u2022 ")):
            text = re.sub(r"^[\s]*[-*\u2022]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_formatted_text(paragraph, text: str):
    """Add text to a paragraph with bold/italic markdown formatting."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            sub_parts = re.split(r"(\*.*?\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and not sub.startswith("**"):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(_strip_markdown(sub))
